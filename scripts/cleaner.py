import docx
import sys
import os

def clean_doc(src_path, output_path):
    print(f"Loading template: {src_path}")
    doc = docx.Document(src_path)
    
    # Strategy: 
    # 1. Identify main tables and wipe content cells (while keeping headers).
    # 2. Identify "Answer Key" sections and wipe them.
    
    # --- Table Cleaning ---
    if len(doc.tables) > 0:
        # Assume Table 0 might be student info. Wipe value cells.
        t0 = doc.tables[0]
        for row in t0.rows:
            # Heuristic: Wipe even index columns (0-based: 1, 3...) if they exist
            # Or wipe based on header text.
            # Let's use a simpler heuristic: if row has >1 cells, wipe the last one?
            # Or distinct key/value pairs.
            for i, cell in enumerate(row.cells):
                # If previous cell looks like a label (short text), wipe this one?
                # Let's rely on standard logic: Wipe all cells except those clearly headers.
                pass
            # For now, let's just wipe specific known patterns if this is a "Lesson Plan".

    # General Table Cleaning (Targeting the big lesson table)
    # Heuristic: Find the table with the most rows.
    target_table = None
    max_rows = 0
    for t in doc.tables:
        if len(t.rows) > max_rows:
            max_rows = len(t.rows)
            target_table = t
            
    if target_table and max_rows > 5:
        print(f"Cleaning main table with {max_rows} rows.")
        for row in target_table.rows:
            # If row has multiple cells, likely Col 0 is Header, Col 1 is Content.
            if len(row.cells) > 1:
                # Determine if it's a section header row or content row.
                # Section headers usually span or have distinct text.
                # If Col 1 has significant text, wipe it.
                row.cells[1].text = "" # Wipe content column
                
                # Check for "Answer" rows and wipe completely?
                if "答案" in row.cells[0].text:
                    row.cells[1].text = ""

    # --- Appendices Cleaning ---
    # Wipe content in paragraphs following "Answer Key" or similar headers
    wipe_mode = False
    for p in doc.paragraphs:
        if "Answer Key" in p.text or "参考答案" in p.text:
            wipe_mode = True
            continue # Keep the header
            
        if wipe_mode:
            p.text = "" # Wipe content

    doc.save(output_path)
    print(f"Cleaned doc saved to {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python cleaner.py <src_docx> <output_docx>")
        sys.exit(1)
        
    src = sys.argv[1]
    dst = sys.argv[2]
    
    clean_doc(src, dst)
