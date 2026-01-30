from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_plan():
    doc = Document()
    
    # --- Styles Setup ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("七年级英语寒假+下学期阶段成长规划")
    run.font.size = Pt(16)
    run.bold = True

    # --- Student Info Table ---
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    # Fill Headers
    info_data = [
        ("学生姓名", "翟以乐", "在读学校", "中科院"), # Assuming school from parsed text
        ("任课老师", "陈老师", "班主任", "许老师"),
        ("入学时间", "2026/1/27", "当前成绩", "55"),
        ("规划阶段", "七年级寒假+下学期", "培养方向", "基础巩固")
    ]
    
    for r, row_data in enumerate(info_data):
        row = table.rows[r]
        for c, text in enumerate(row_data):
            row.cells[c].text = text
            # Center align
            row.cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n") # Spacer

    # --- Analysis Section ---
    h1 = doc.add_heading('一、学情深度分析 (Student Analysis)', level=1)
    
    p = doc.add_paragraph()
    p.add_run("1. 性格特点：").bold = True
    p.add_run("活泼开朗，思维活跃，但关注力（Focus）较难持久，需在课堂中增加互动频次，采用“短时多次”的知识输出方式。\n")
    
    p.add_run("2. 现有基础：").bold = True
    p.add_run("当前分数55分（基础薄弱）。主要问题集中在词汇量不足（小学核心词汇遗忘）和基础语法框架（词性、时态）缺失。阅读理解常靠“猜”，缺乏逻辑支撑。\n")
    
    p.add_run("3. 提升策略：").bold = True
    p.add_run("采用“陪读+背默+日练”三维一体模式。课上重构语法体系，课下死磕词汇过关。")

    doc.add_paragraph("\n")

    # --- 20-Lesson Roadmap Table ---
    h2 = doc.add_heading('二、20次课程详细规划 (Review & Advance)', level=1)
    
    # Table Header
    plan_table = doc.add_table(rows=21, cols=4)
    plan_table.style = 'Table Grid'
    plan_table.autofit = False 
    plan_table.allow_autofit = False
    
    headers = ["次序", "模块", "核心内容 (Core Layout)", "难度"]
    for i, h in enumerate(headers):
        cell = plan_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True

    # Plan Data (Designed for Grade 7 Basics)
    lessons = [
        ("1", "词法基础", "名词(单复数/所有格) + 冠词(a/an/the)逻辑", "易"),
        ("2", "词法基础", "代词全家桶 (主/宾/形物/名物/反身)", "易"),
        ("3", "词法基础", "数词表达 (时间/年代/分数) + 介词初探", "易"),
        ("4", "句法基础", "Be动词 vs 实义动词 (否定/疑问转换)", "中"),
        ("5", "句法基础", "There be 句型与就近原则", "中"),
        ("6", "时态进阶", "一般现在时 (三单变化规则)", "难"),
        ("7", "时态进阶", "现在进行时 (ing变化规则)", "中"),
        ("8", "时态进阶", "一般过去时 (ed与不规则动词表)", "难"),
        ("9", "时态进阶", "一般将来时 (will vs be going to)", "中"),
        ("10", "阶段测评", "前9次课综合复习 + 阶段测验1", "难"),
        ("11", "句型突破", "祈使句与感叹句 (What/How)", "易"),
        ("12", "句型突破", "情态动词 (can/must/may) 用法", "中"),
        ("13", "词汇专项", "七年级核心动词短语搭配 (Get/Take/Have)", "难"),
        ("14", "阅读技巧", "阅读理解A篇：信息提取与细节定位", "中"),
        ("15", "阅读技巧", "完形填空：上下文逻辑与词义辨析", "难"),
        ("16", "综合语法", "形容词与副词 (比较级/最高级)", "难"),
        ("17", "综合语法", "简单句的五大基本句型分析", "难"),
        ("18", "写作专项", "基础写作：句子翻译与看图写话", "中"),
        ("19", "总复习", "期末重难点易错题刷题", "难"),
        ("20", "期末冲刺", "期末全真模拟考 + 试卷分析", "难")
    ]

    for i, (idx, mod, content, diff) in enumerate(lessons):
        row = plan_table.rows[i+1]
        row.cells[0].text = f"第{idx}次"
        row.cells[1].text = mod
        row.cells[2].text = content
        row.cells[3].text = diff
        
        # Center align first and last cols
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")
    
    # --- Footer/Sign ---
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.add_run("规划师/教师：").bold = True
    p_sign.add_run("陈老师   ")
    p_sign.add_run("日期：").bold = True
    p_sign.add_run("2026/01/30")

    doc.save("output_test/Zhai_Yile_Plan_v1.docx")
    print("Plan saved to output_test/Zhai_Yile_Plan_v1.docx")

if __name__ == "__main__":
    create_plan()
