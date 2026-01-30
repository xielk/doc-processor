import docx
import sys

def inspect(path):
    doc = docx.Document(path)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    for i, t in enumerate(doc.tables):
        print(f"Table {i}: {len(t.rows)} rows, {len(t.columns)} cols")
        if len(t.rows) > 0:
            print(f"  Row 0 sample: {[c.text.strip() for c in t.rows[0].cells]}")

if __name__ == "__main__":
    inspect(sys.argv[1])
