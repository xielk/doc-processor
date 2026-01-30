#!/usr/bin/env python3
"""
Resource Indexer for Question Bank
生成题库索引，支持按需加载
"""

import os
import json
import glob
from docx import Document
from pathlib import Path
from datetime import datetime

RESOURCE_PATH = "/Users/xielk/webdata/english/lesson/resource"
INDEX_FILE = os.path.join(RESOURCE_PATH, "index.json")

def extract_metadata_from_path(file_path):
    """从文件路径解析元数据"""
    parts = Path(file_path).parts
    
    # 尝试提取年份
    year = "unknown"
    for part in parts:
        if "202" in part or "20" in part:
            if "2025" in part:
                year = "2025"
            elif "2024" in part:
                year = "2024"
            elif "2023" in part:
                year = "2023"
            break
    
    # 尝试提取区域
    districts = ["徐汇", "浦东", "嘉定", "黄浦", "静安", "虹口", "杨浦", "长宁", "普陀", "宝山", "闵行", "松江", "金山", "青浦", "奉贤", "崇明", "上海"]
    district = "unknown"
    for part in parts:
        for d in districts:
            if d in part:
                district = d
                break
        if district != "unknown":
            break
    
    # 尝试提取考试类型
    exam_types = ["一模", "二模", "中考", "期末", "期中"]
    exam_type = "unknown"
    for part in parts:
        for et in exam_types:
            if et in part:
                exam_type = et
                break
        if exam_type != "unknown":
            break
    
    # 尝试提取题型
    question_types = {
        "语法": ["语法", "非谓语", "从句", "时态", "语态"],
        "阅读": ["阅读", "A篇", "B篇", "C篇", "D篇", "完形"],
        "作文": ["作文", "写作", "范文"],
        "词汇": ["词汇", "单词", "短语"],
        "听力": ["听力", "听说"],
        "综合": ["综合", "模拟", "真题"]
    }
    
    q_type = "综合"
    file_lower = file_path.lower()
    for qt, keywords in question_types.items():
        for kw in keywords:
            if kw in file_lower or kw in str(parts):
                q_type = qt
                break
        if q_type != "综合":
            break
    
    return {
        "year": year,
        "district": district,
        "exam_type": exam_type,
        "question_type": q_type
    }

def extract_preview(docx_path, max_chars=500):
    """提取文档预览内容"""
    try:
        doc = Document(docx_path)
        preview_parts = []
        char_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                preview_parts.append(text)
                char_count += len(text)
                if char_count >= max_chars:
                    break
        
        # 也尝试提取表格中的文本（前3个表格）
        for table in doc.tables[:3]:
            for row in table.rows[:3]:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in preview_parts:
                        preview_parts.append(text)
                        char_count += len(text)
                        if char_count >= max_chars:
                            break
                if char_count >= max_chars:
                    break
            if char_count >= max_chars:
                break
        
        return " ".join(preview_parts)[:max_chars]
    except Exception as e:
        print(f"Error extracting preview from {docx_path}: {e}")
        return ""

def create_index():
    """创建题库索引"""
    print("开始创建题库索引...")
    print(f"搜索路径: {RESOURCE_PATH}")
    
    index = []
    total_size = 0
    
    # 查找所有docx文件
    docx_files = glob.glob(os.path.join(RESOURCE_PATH, "**/*.docx"), recursive=True)
    print(f"找到 {len(docx_files)} 个docx文件")
    
    for i, docx_file in enumerate(docx_files, 1):
        try:
            # 提取元数据
            metadata = extract_metadata_from_path(docx_file)
            
            # 提取预览
            preview = extract_preview(docx_file)
            
            # 获取文件信息
            stat = os.stat(docx_file)
            file_info = {
                "id": i,
                "file": docx_file,
                "filename": os.path.basename(docx_file),
                "year": metadata["year"],
                "district": metadata["district"],
                "exam_type": metadata["exam_type"],
                "question_type": metadata["question_type"],
                "preview": preview,
                "size_kb": stat.st_size // 1024,
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d")
            }
            
            index.append(file_info)
            total_size += stat.st_size
            
            if i % 50 == 0:
                print(f"已处理 {i}/{len(docx_files)} 个文件...")
                
        except Exception as e:
            print(f"处理文件出错 {docx_file}: {e}")
            continue
    
    # 保存索引
    index_data = {
        "metadata": {
            "total_files": len(index),
            "total_size_mb": total_size / (1024 * 1024),
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "resource_path": RESOURCE_PATH
        },
        "files": index
    }
    
    with open(INDEX_FILE, 'w', encoding='utf-8') as f:
        json.dump(index_data, f, ensure_ascii=False, indent=2)
    
    print(f"\n✅ 索引创建完成!")
    print(f"   - 文件总数: {len(index)}")
    print(f"   - 总大小: {total_size / (1024 * 1024):.2f} MB")
    print(f"   - 索引位置: {INDEX_FILE}")
    
    # 显示统计信息
    print("\n📊 统计信息:")
    years = {}
    districts = {}
    q_types = {}
    
    for item in index:
        years[item["year"]] = years.get(item["year"], 0) + 1
        districts[item["district"]] = districts.get(item["district"], 0) + 1
        q_types[item["question_type"]] = q_types.get(item["question_type"], 0) + 1
    
    print(f"   年份分布: {dict(sorted(years.items()))}")
    print(f"   区域分布: {dict(sorted(districts.items()))}")
    print(f"   题型分布: {dict(sorted(q_types.items()))}")
    
    return index_data

if __name__ == "__main__":
    create_index()
